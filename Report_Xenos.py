#!/usr/bin/env python

# ------------------------------------------------------------------------------
# 0) DOWNLOAD THE REQUIRED PACKAGES
# ------------------------------------------------------------------------------
from webbot import Browser

import time

import os

import pandas as pd
import numpy as np

from lib import custom_cleaning as c_cleaning
from lib import custom_aggregate as c_aggregate
from lib import custom_graphs as c_graphs
from lib import document_styles
from lib import to_pdf

from docx.shared import Pt, Cm

# ------------------------------------------------------------------------------
# 1) SOURCE CSV DOWNLOAD
# ------------------------------------------------------------------------------
# Open the default browser and go to the website
web = Browser()
web.go_to("https://my.bigcartel.com")

# Provide login information and log in
web.type("xenosmilan.**********",
         into="account_subdomain",
         id="account_subdomain")
web.type("****", into="password", id="password")
web.click("Log In")

# Navigate through the website
web.click("Orders")
web.click("Shipped")

# Downloadthereport
web.go_to("https://****/orders_exports.csv")
time.sleep(5)

# Close the browser
web.quit()

# NOTE##########################################
# Setting the directory into the resources folder
# ##############################################
resources_path = os.getcwd() + "/Resources"

os.chdir(resources_path)
# ------------------------------------------------------------------------------
# 2) INDENTIFY SOURCE CSV THROUGH THE PATH IN THE TXT FILE
# ------------------------------------------------------------------------------

with open('Report_path.txt', 'r') as file:
    data = file.read()
download_dir = data.split("\n")[0].replace('"', "").replace("\\", "/")

# ------------------------------------------------------------------------------
# 3) MOVING THE SOURCE CSV TO THE RESOURCES FOLDER
# ------------------------------------------------------------------------------
# Get the download directory and file name
filename = "orders.csv"

# Indentify location of origin and destination inside the directory
old_report_path = download_dir + "/" + filename
new_report_path = resources_path + "/" + filename

# Move the file
os.replace(old_report_path, new_report_path)

# ------------------------------------------------------------------------------
# 4) REPORT ANALYSIS
# ------------------------------------------------------------------------------
# ### 4.0) Getting the table ready
# Open the table
df = pd.read_csv(new_report_path)

# ------------------------------------------------------------------------------
# 4) REPORT ANALYSIS
# ------------------------------------------------------------------------------
# ### 4.0) Getting the table ready
# Open the table
df = pd.read_csv("orders.csv")
# FROM HERE ON

# Delete useless columns
to_drop = ["Buyer email", "Buyer phone number", "Transaction ID", "Time",
           "Shipping address 1", "Shipping address 2", "Currency",
           "Total tax", "Tax remitted by Big Cartel", "Note",
           "Private notes"]
df.drop(to_drop, axis=1, inplace=True)

# Unpack rows containing different items
df = c_cleaning.unpack_multiple_orders(df)

# Estimate gender based on the name of the customer, creating a boolean column
df["Male"] = c_cleaning.estimate_gender(df, "Buyer first name")

# Derive city, province and region from zip code
df = c_cleaning.match_zip_to_city(df)

# Aggregate the three status columns into a single one
to_aggregate = ["Status", "Payment status", "Shipping status"]
df["Order competed"] = c_cleaning.aggregate_status(df, to_aggregate, drop=True)

# Convert date into datetime format
df["Date"] = pd.to_datetime(df["Date"], yearfirst=True)

# Add a column for a univoque monthcode
df["Month code"] = c_cleaning.gen_month_code(df, "Date")

# Create single column for full name
df["Name surname"] = df["Buyer first name"] + " " + df["Buyer last name"]

# Extracting the size from the items columns
df["Size"] = df["Items"].str.extract(pat="(on_name:[A-Z]?[A-Z]?[A-Z])")
df["Size"] = df["Size"].str.slice(start=8)

# Exctracting the item name from the items column
df["Items"] = df["Items"].str.split(pat="|").str[0].str.split(pat=":").str[1]

# Adding the item type from an external register
dr_items = pd.read_excel("Info_capi.xlsx", header=0)
df = pd.merge(df, dr_items[["Name", "Items type"]],
              left_on="Items", right_on="Name", how="left")

# Define net earnings for each item
df["Net earnings"] = df["Item total"] - df["Total discount"]

# Define new column names and apply them
new_names = {"Number" : "Code",
             "Item count" : "Items count",
             "Item total" : "Raw price",
             "Total price" : "Paid price",
             "Total shipping" : "Shipping price",
             "Total discount" : "Discount"}
df.rename(new_names, axis=1, inplace=True)

# Rule out the columns currently considered useless
df = df[["Code", "Name surname", "Male", "Month code", "Date", "Order competed",
         "City", "Province", "Region", "Items", "Items type", "Size",
         "Items count", "Raw price", "Paid price", "Shipping price",
         "Discount", "Net earnings"]]

df.sort_values(by=["Date"], inplace=True, ascending=False)

# Saving the formatted report, overwriting if already present
filename = "Orders updated.xlsx"
if os.path.isfile(filename): os.remove(filename)
df.to_excel(filename, index=False)

# ------------------------------------------
# ### 4.2) Creating report for total values
# Generate descriptive indicators for volume of sales and revenues

values_tot = {}
values_mean = {}
values_min = {}
values_max = {}

columns_of_interest = ["Items count", "Raw price", "Net earnings"]

for column in columns_of_interest:
    values_tot[column] = round(df[column].sum(), 0)
    values_mean[column] = round(df[column].mean(), 1)
    values_min[column] = round(df[column].min(), 0)
    values_max[column] = round(df[column].max(), 0)

df1 = c_aggregate.aggregate_by_date(df, "Month code", "Month")

c_graphs.plot_vert_time_serie(df1, "Month", "Tot items ordered",
                              "Tot net earnings",
                              "Males percentage",
                              title="Items sold and revenues over time",
                              x1_title="Months",
                              y1_title="Item counts",
                              y2_title="Net earnings",
                              save=True,
                              file_name="Items sold and revenues_time.png",
                              show=False)

c_graphs.plot_cumulative_time_serie(df1, "Month", "Cumulative items sold",
                                    "Cumulative net earnings",
                                    "Cumulative sales and earnings", "Month",
                                    "Items sold", "Net earnings",
                                    save=True,
                                    file_name="Cumulative sales and earnings_time.png",
                                    show=False)

df2 = c_aggregate.aggregate_by_category(df, "Items type", "Items type")

c_graphs.plot_horizontal_bar(dataframe=df2,
                             y1_serie="Items type",
                             x1_serie="Tot items count",
                             x2_serie="Tot net earnings",
                             gender_serie="Males percentage",
                             title="Indicators for item type",
                             y1_title="Item type",
                             x1_title="Item count",
                             x2_title="Net earnings",
                             save=True,
                             file_name="Indicators for item type.png",
                             show=False)

df3 = c_aggregate.aggregate_by_category(df, "Items", "Items")

c_graphs.plot_horizontal_bar(dataframe=df3,
                             y1_serie="Items",
                             x1_serie="Tot items count",
                             x2_serie="Tot net earnings",
                             gender_serie="Males percentage",
                             title="Indicators for item",
                             y1_title="Item",
                             x1_title="Item count",
                             x2_title="Net earnings",
                             save=True,
                             file_name="Indicators for item.png",
                             show=False)

# ------------------------------------------------------------------------------
# 5) Document setting and customization
# ------------------------------------------------------------------------------
# ### Document setting and customization
document = document_styles.style_document_xenos("Xenos report template.docx")

# ------------------------------------------------------------------------------
# 6) Writing in the document
# ------------------------------------------------------------------------------

title = document.add_paragraph(str.upper("Report Xenos"))
title.style = document.styles["Custom title"]

paragraph2 = document.add_paragraph("Dati storici")
paragraph2.style = document.styles["Custom heading 1"]

text = ("Questa sezione contiene i dati calcolati su tutto il periodo di attività.")
paragraph3 = document.add_paragraph(text)
paragraph3.style = document.styles["Custom body"]

paragraph4 = document.add_paragraph("Overview")
paragraph4.style = document.styles["Custom heading 2"]

text1 = f"Numero totale di capi venduti: {values_tot['Items count']}"
paragraph5 = document.add_paragraph(text1)
paragraph5.style = document.styles["Custom body"]

text2 = f"Numero medio di capi venduti per ordine: {values_mean['Items count']}"
paragraph6 = document.add_paragraph(text2)
paragraph6.style = document.styles["Custom body"]

text3 = f"Numero massimo di capi venduti in un ordine: {values_max['Items count']}"
paragraph7 = document.add_paragraph(text3)
paragraph7.style = document.styles["Custom body"]

text4 = f"Valore complessivo dei capi acquistati: {values_tot['Raw price']}"
paragraph6 = document.add_paragraph(text4)
paragraph6.style = document.styles["Custom body"]

text5 = f"Valore medio dei capi acquistati per ordine: {values_mean['Raw price']}"
paragraph8 = document.add_paragraph(text5)
paragraph8.style = document.styles["Custom body"]

text6 = f"Valore massimo dei capi acquistati in un ordine: {values_mean['Raw price']}"
paragraph9 = document.add_paragraph(text6)
paragraph9.style = document.styles["Custom body"]

text7 = f"Totale ricavi al netto dei costi di vendita: {values_tot['Net earnings']}"
paragraph10 = document.add_paragraph(text7)
paragraph10.style = document.styles["Custom body"]

text8 = f"Media ricavi al netto dei costi di vendita per ordine: {values_mean['Net earnings']}"
paragraph11 = document.add_paragraph(text8)
paragraph11.style = document.styles["Custom body"]

text9 = f"Massimo ricavi al netto dei costi di vendita in un ordine: {values_max['Net earnings']}"
paragraph12 = document.add_paragraph(text9)
paragraph12.style = document.styles["Custom body"]

paragraph13 = document.add_paragraph("Dati nel tempo")
paragraph13.style = document.styles["Custom heading 2"]

text = "Il seguente grafico mostra l'andamento del numero di vendite e dei\
 ricavi al netto dei costi di vendita nel corso del tempo. L'unità di misura\
 utilizzata è il mese."
paragraph14 = document.add_paragraph(text)
paragraph14.style = document.styles["Custom body"]
document.add_picture("Items sold and revenues_time.png", width=Cm(15.0))

text = "Il grafico successivo mostra invece l'andamento del tempo della\
 quantità cumulativa di numeri di capi venduti e ricavi al netto dei costi\
 di vendita. La quantità cumulativa è ottenuta sommando il valore del periodo\
 alla somma di tutti i valori precedenti."
paragraph15 = document.add_paragraph(text)
paragraph15.style = document.styles["Custom body"]
document.add_picture("Cumulative sales and earnings_time.png",
                     width=Cm(15.0))

paragraph16 = document.add_paragraph("Dati per tipologia di capo")
paragraph16.style = document.styles["Custom heading 2"]

text = "Il grafico che segue mostra il numero totale di ordini per tipologia\
 di capo, divisa in due colorazioni a seconda del genere. Insieme a questo,\
 sul secondo asse, sono indicati i ricavi totali al netto dei costi di vendita\
 per tipologia di prodotto."
paragraph17 = document.add_paragraph(text)
paragraph17.style = document.styles["Custom body"]
document.add_picture("Indicators for item type.png",
                     width=Cm(15.0))

paragraph17 = document.add_paragraph("Dati per capo")
paragraph17.style = document.styles["Custom heading 2"]

text = "Il grafico è del tutto simile al precedente, con i singoli capi al\
 posto delle tipologie di capo."
paragraph18 = document.add_paragraph(text)
paragraph18.style = document.styles["Custom body"]
document.add_picture("Indicators for item.png",
                     width=Cm(15.0))

filename = "Xenos report.docx"
document.save(filename)

# ------------------------------------------------------------------------------
# 7) Transform the document into a pdf
# ------------------------------------------------------------------------------

path = os.getcwd() + "/" + filename
to_pdf.convert_to_pdf(path)

# ------------------------------------------------------------------------------
# 8) Move the document to the main folder
# ------------------------------------------------------------------------------

pdf_filename = filename.replace(".docx", ".pdf")
old_pdf_path = os.getcwd() + "/" + pdf_filename
new_pdf_path = os.getcwd().replace("Resources", pdf_filename)
os.replace(old_pdf_path, new_pdf_path)

print("Process completed")
input("Press any key to exit")
