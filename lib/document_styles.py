# -*- coding: utf-8 -*-
"""
Created on Wed Jun  3 18:36:45 2020

@author: Filippo
"""


def style_document_xenos(template_name):

    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
    from datetime import date

    document = Document(template_name)

    # Footer
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    today = date.today().strftime("%d/%m/%Y")
    run = paragraph.add_run(f"\tConfidenziale - Xenos Milan\t{today}")
    run.italic = True

    # ###Define the various styles in descending hierarchical order
    my_styles = document.styles

    # Title
    title_style = my_styles.add_style("Custom title", WD_STYLE_TYPE.PARAGRAPH)
    title_style.base_style = my_styles['Normal']
    format_title_style = title_style.paragraph_format
    title_style.hidden = False
    title_style.quick_style = True
    title_style.priority = 2

    # Heading 1
    heading_1_style = my_styles.add_style("Custom heading 1",
                                          WD_STYLE_TYPE.PARAGRAPH)
    heading_1_style.base_style = my_styles['Normal']
    format_heading_1_style = heading_1_style.paragraph_format
    heading_1_style.hidden = False
    heading_1_style.quick_style = True
    heading_1_style.priority = 3

    # Heading 2
    heading_2_style = my_styles.add_style("Custom heading 2",
                                          WD_STYLE_TYPE.PARAGRAPH)
    heading_2_style.base_style = my_styles['Normal']
    format_heading_2_style = heading_2_style.paragraph_format
    heading_2_style.hidden = False
    heading_2_style.quick_style = True
    heading_2_style.priority = 4

    # Heading 3
    heading_3_style = my_styles.add_style("Custom heading 3",
                                          WD_STYLE_TYPE.PARAGRAPH)
    heading_3_style.base_style = my_styles['Normal']
    format_heading_3_style = heading_3_style.paragraph_format
    heading_3_style.hidden = False
    heading_3_style.quick_style = True
    heading_3_style.priority = 5

    # Body
    body_style = my_styles.add_style("Custom body",
                                     WD_STYLE_TYPE.PARAGRAPH)
    body_style.base_style = my_styles['Normal']
    format_body_style = body_style.paragraph_format
    body_style.hidden = False
    body_style.quick_style = True
    body_style.priority = 1

    # ### Customization
    # Customize title
    format_title_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_title_style.space_before = Pt(3.0)
    format_title_style.space_after = Pt(12.0)
    format_title_style.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    title_style.font.name = "Arial Black"
    title_style.font.bold = True
    title_style.font.underline = WD_UNDERLINE.THICK
    title_style.font.size = Pt(20)

    # Customize heading 1
    format_heading_1_style.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_heading_1_style.space_before = Pt(12.0)
    format_heading_1_style.space_after = Pt(12.0)
    format_heading_1_style.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading_1_style.font.name = "Arial"
    heading_1_style.font.italic = True
    heading_1_style.font.bold = True
    heading_1_style.font.underline = True
    heading_1_style.font.size = Pt(16)

    # Customize heading 2
    format_heading_2_style.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_heading_2_style.space_before = Pt(8.0)
    format_heading_2_style.space_after = Pt(6.0)
    format_heading_2_style.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading_2_style.font.name = "Arial"
    heading_2_style.font.italic = True
    heading_2_style.font.bold = True
    heading_2_style.font.underline = True
    heading_2_style.font.size = Pt(13)

    # Customize heading 3
    format_heading_3_style.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_heading_3_style.space_before = Pt(6.0)
    format_heading_3_style.space_after = Pt(6.0)
    format_heading_3_style.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading_3_style.font.name = "Arial"
    heading_3_style.font.italic = True
    heading_3_style.font.bold = True
    heading_3_style.font.size = Pt(11)

    # Customize body
    format_body_style.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    format_body_style.space_before = Pt(3.0)
    format_body_style.space_after = Pt(3.0)
    format_body_style.line_spacing_rule = WD_LINE_SPACING.SINGLE

    body_style.font.name = "Arial"
    body_style.font.size = Pt(11)

    return document
